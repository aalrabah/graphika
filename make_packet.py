#!/usr/bin/env python3
r"""Generate an instructor review packet (.docx) from a course's relations.jsonl.

Reproduces the approved template (InstructKG Concept Map Review, SAMPLE
cs401_403); only course-specific content changes between packets.

Set the course once, then run:

    C=cs401_403
    TITLE="iCAN Algorithms (CS 401/403)"
    M=qwen14b

    python make_packet.py \
        --relations out/$C/relations.jsonl \
        --course-title "$TITLE" \
        --run-label "$C, Qwen-14B, August 2026" \
        --eval-json out/$C/eval_${C}_${M}.json \
        --out out/$C/InstructKG_Review_Packet_${C}.docx

If there is no eval JSON, drop --eval-json and pass --node-significance and
--triplet-accuracy from the eval log instead.

Sampling is deterministic for a given --seed (default 597): Part A takes the 12
most-referenced concepts plus 3 seeded draws from the tail (unfiltered); Part B
is a seeded random sample of 15 real (non-null) edges.

Requires python-docx. Written by Claude (Anthropic) for Jeffry Royce, 2026-08.
"""

import argparse
import json
import random
import sys
from collections import Counter

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------- palette
TITLE_BLUE = RGBColor(0x00, 0x55, 0xC7)
HEADING_BLUE = RGBColor(0x4A, 0x86, 0xE8)
SAMPLE_RED = RGBColor(0xFF, 0x00, 0x00)
BODY_GRAY = RGBColor(0x22, 0x22, 0x22)
META_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_NAVY = "1F4E79"     # Part A/B table header fill
LABEL_GRAY = "F2F2F2"      # page-1 table label fill
BORDER_GRAY = "BFBFBF"

REL_WORD = {"depends_on": "DEPENDS ON", "part_of": "PART OF"}
REL_WORD_GLANCE = {"depends_on": "depends-on", "part_of": "part-of"}

# ---------------------------------------------------------------- helpers


def prettify(name, overrides):
    if name in overrides:
        return overrides[name]
    s = name.replace("_", " ").strip().lower()
    return s[:1].upper() + s[1:] if s else name


def load_relations(path):
    records = []
    with open(path) as f:
        for i, line in enumerate(f):
            line = line.strip()
            if not line:
                continue
            try:
                records.append(json.loads(line))
            except json.JSONDecodeError as e:
                print(f"  ! skipping malformed line {i + 1}: {e}", file=sys.stderr)
    return records


def evidence_count(rec):
    info = (rec.get("_meta") or {}).get("mode_info") or {}
    return (
        info.get("chunk_co_occurrence_count")
        or info.get("cluster_co_occurrence_count")
        or len(rec.get("evidence_chunks") or [])
    )


# ---------------------------------------------------------------- docx style


def set_cell_bg(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_borders(table, hex_color=BORDER_GRAY):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def set_col_widths(table, widths):
    table.autofit = False
    # rewrite tblGrid so LibreOffice/Google Docs honor the widths too
    grid = table._tbl.tblGrid
    for gc in list(grid):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.emu / 635)))  # EMU -> twips
        grid.append(gc)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def run(paragraph, text, size=11, bold=False, italic=False, color=BODY_GRAY,
        font="Arial"):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def para(doc, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    return p


def heading(doc, text, before=18, after=6, new_page=False):
    p = para(doc, before=before, after=after)
    if new_page:
        p.paragraph_format.page_break_before = True
    run(p, text, size=20, color=HEADING_BLUE)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run(p, bold_prefix, bold=True)
    run(p, text)
    return p


def cell_para(cell):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def answer_rule(doc):
    """Blank answer line rendered as a paragraph bottom border."""
    p = para(doc, before=18, after=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------- content


def build(args):
    records = load_relations(args.relations)
    if not records:
        sys.exit(f"No records loaded from {args.relations}")

    overrides = {}
    if args.name_overrides:
        with open(args.name_overrides) as f:
            overrides = json.load(f)

    rng = random.Random(args.seed)

    real = [r for r in records if r.get("relation")]
    all_names = {r[s]["name"] for r in records for s in ("A", "B")}
    edge_names = {r[s]["name"] for r in real for s in ("A", "B")}
    rel_counts = Counter(r["relation"] for r in real)

    print(f"records={len(records)} real_edges={len(real)} "
          f"concepts(all records)={len(all_names)} concepts(real edges only)={len(edge_names)}")
    print(f"relation mix: {dict(rel_counts)}")

    # frequency over real edges (degree with multiplicity)
    freq = Counter()
    for r in real:
        freq[r["A"]["name"]] += 1
        freq[r["B"]["name"]] += 1
    ranked = sorted(freq, key=lambda n: (-freq[n], n))

    head = ranked[: args.head_concepts]
    tail_pool = ranked[args.head_concepts:]
    tail = (rng.sample(tail_pool, min(args.tail_concepts, len(tail_pool)))
            if tail_pool else [])
    part_a = head + tail

    part_b = rng.sample(real, min(args.relationships, len(real)))

    top_edges = sorted(real, key=evidence_count, reverse=True)
    examples, seen = [], set()
    for e in top_edges:
        pair = frozenset((e["A"]["name"], e["B"]["name"]))
        if pair in seen:
            continue
        seen.add(pair)
        examples.append(e)
        if len(examples) == 2:
            break

    scores = (args.node_significance, args.triplet_accuracy)
    if args.eval_json:
        with open(args.eval_json) as f:
            ev = json.load(f)
        scores = (ev["node_significance"]["mean"], ev["triplet_accuracy"]["mean"])
    if scores[0] is None or scores[1] is None:
        sys.exit("Provide --eval-json or both --node-significance and --triplet-accuracy")

    # ------------------------------------------------------------ document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_GRAY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    usable = Inches(6.5)

    # title
    p = para(doc, after=6)
    run(p, "InstructKG Concept Map Review — ", size=26, bold=True, color=TITLE_BLUE)
    run(p, args.banner_label, size=26, bold=True,
        color=SAMPLE_RED if args.sample else TITLE_BLUE)

    for text in (f"Course: {args.course_title}", f"Run: {args.run_label}"):
        p = para(doc, after=2)
        run(p, text, color=META_GRAY)
    p = para(doc, before=8, after=2)
    run(p, "Prepared by Jeffry Royce (CS 597, directed study) "
           "with Prof. Abdussalam Alawini", color=META_GRAY)
    p = para(doc, after=2)
    run(p, "Built on InstructKG (AlRabah, Kargupta, Han, Alawini)", color=META_GRAY)

    # What this is
    heading(doc, "What this is", before=24)
    p = para(doc, after=6)
    run(p, "InstructKG read this course's lecture materials and inferred a "
           "concept map: the concepts taught, and how they depend on each "
           "other. This packet shows a sample of what it produced and "
           "requests ~20–30 minutes of instructor judgment regarding quick "
           "0/1/2 ratings of a small sample of concepts and relationships.")

    # LLM-judged scores
    heading(doc, "LLM-judged scores for this graph:")
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    set_col_widths(t, [Inches(3.25), Inches(3.25)])
    for j, text in enumerate(("Metric", "Score")):
        c = t.rows[0].cells[j]
        set_cell_bg(c, LABEL_GRAY)
        run(cell_para(c), text, bold=True)
    for i, (label, val) in enumerate((
            ("Node significance (concepts)", scores[0]),
            ("Triplet accuracy (relationships)", scores[1])), start=1):
        run(cell_para(t.rows[i].cells[0]), label)
        run(cell_para(t.rows[i].cells[1]), f"{val:.2f}")

    p = para(doc, before=10, after=6)
    run(p, "The scores above represent the AI model scoring itself. Concepts "
           "scores close to 1.00 indicate nearly everything extracted looks "
           "like real course content. A relationship score of 0.75 indicates "
           "a quarter of the proposed relationships look incorrect, "
           "reversed, or debatable.")

    # This course at a glance
    heading(doc, "This course at a glance")
    most_ref = ", ".join(prettify(n, overrides) for n in head[: args.glance_concepts])
    rel_break = ", ".join(f"{REL_WORD_GLANCE[k]}: {rel_counts[k]}"
                          for k in ("depends_on", "part_of") if k in rel_counts)
    ex_links = " · ".join(
        f"{prettify(e['A']['name'], overrides)} {REL_WORD[e['relation']]} "
        f"{prettify(e['B']['name'], overrides)}" for e in examples)
    glance = [
        ("Concepts extracted", f"{len(all_names)} unique concepts"),
        ("Relationships inferred",
         f"{len(real)} unique typed relationships ({rel_break})"),
        ("Most-referenced concepts", most_ref),
        ("Example inferred links", ex_links),
    ]
    t = doc.add_table(rows=len(glance), cols=2)
    set_table_borders(t)
    set_col_widths(t, [Inches(3.25), Inches(3.25)])
    for i, (label, val) in enumerate(glance):
        c = t.rows[i].cells[0]
        set_cell_bg(c, LABEL_GRAY)
        run(cell_para(c), label, bold=True)
        run(cell_para(t.rows[i].cells[1]), val)

    # How to read the relationships
    heading(doc, "How to read the relationships", before=0, new_page=True)
    p = para(doc, after=8)
    run(p, "A ")
    run(p, "DEPENDS ON", bold=True)
    run(p, " B: to understand A, a student should first understand B (B is a "
           "prerequisite of A).")
    p = para(doc, after=8)
    run(p, "Example: “Merge sort DEPENDS ON Recursion.”")
    p = para(doc, after=8)
    run(p, "A ")
    run(p, "PART OF", bold=True)
    run(p, " B: A is a component, subtype, or member of the broader concept B.")
    p = para(doc, after=8)
    run(p, "Example: “Depth-first search PART OF Graph traversal.”")
    p = para(doc, before=6, after=6)
    run(p, "Direction matters. Related concepts with the arrow pointing the "
           "wrong way (or the wrong type) score a 1.")

    # Part A
    heading(doc, "Part A — Are these significant concepts in this course?")
    p = para(doc, after=6)
    run(p, "“Is this a meaningful piece of course content students should be "
           "taught and understand?”")
    bullet(doc, " — Yes: clearly valid, significant course concept (core "
                "topic, method, principle, key term).", bold_prefix="2")
    bullet(doc, " — Sort of: plausible but vague, overly broad, or only "
                "weakly part of the course.", bold_prefix="1")
    bullet(doc, " — No: not course content, or clearly insignificant.",
           bold_prefix="0")
    para(doc, after=4)

    t = doc.add_table(rows=len(part_a) + 1, cols=4)
    set_table_borders(t)
    widths = [Inches(0.55), Inches(3.0), Inches(1.2), Inches(1.75)]
    set_col_widths(t, widths)
    for j, text in enumerate(("#", "Concept", "Score (0/1/2)", "Optional note")):
        c = t.rows[0].cells[j]
        set_cell_bg(c, HEADER_NAVY)
        run(cell_para(c), text, bold=True, color=WHITE)
    for i, name in enumerate(part_a, start=1):
        run(cell_para(t.rows[i].cells[0]), str(i))
        run(cell_para(t.rows[i].cells[1]), prettify(name, overrides))

    p = para(doc, before=4, after=0)
    note = (f"Concepts sampled across the frequency range — "
            f"{args.head_concepts} most-referenced plus {args.tail_concepts} "
            f"drawn from the tail, deliberately unfiltered.")
    if args.part_a_note:
        note += " " + args.part_a_note
    run(p, note, size=9, italic=True, color=META_GRAY)

    # Part B
    heading(doc, "Part B — Are these relationships right?", before=0,
            new_page=True)
    p = para(doc, after=4)
    run(p, "Judge two things:")
    for i, text in enumerate((
            "are the concepts directly related as course concepts",
            "are the relation type and direction correct?"), start=1):
        p = para(doc, after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        run(p, f"({i}) {text}")
    bullet(doc, " — Yes: directly related, and type and direction are "
                "correct.", bold_prefix="2")
    bullet(doc, " — Somewhat: related, but the type and/or direction is "
                "wrong or unclear.", bold_prefix="1")
    bullet(doc, " — No: should not be directly related, or unsupported.",
           bold_prefix="0")
    para(doc, after=4)

    t = doc.add_table(rows=len(part_b) + 1, cols=4)
    set_table_borders(t)
    widths = [Inches(0.55), Inches(3.55), Inches(0.95), Inches(1.45)]
    set_col_widths(t, widths)
    for j, text in enumerate(("#", "Proposed relationship (with system's "
                              "justification)", "Score (0/1/2)",
                              "Optional note")):
        c = t.rows[0].cells[j]
        set_cell_bg(c, HEADER_NAVY)
        run(cell_para(c), text, bold=True, color=WHITE)
    for i, rec in enumerate(part_b, start=1):
        run(cell_para(t.rows[i].cells[0]), str(i))
        c = t.rows[i].cells[1]
        p = cell_para(c)
        run(p, f"{prettify(rec['A']['name'], overrides)}   "
               f"{REL_WORD[rec['relation']]}   "
               f"{prettify(rec['B']['name'], overrides)}", bold=True)
        p2 = c.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(2)
        just = (rec.get("justification") or "").strip()
        if len(just) > args.justification_chars:
            just = just[: args.justification_chars].rsplit(" ", 1)[0] + "…"
        run(p2, just, size=9, italic=True, color=META_GRAY)

    p = para(doc, before=4, after=0)
    run(p, "Relationships randomly sampled (seeded) from the run. "
           "Deliberately not cherry-picked, so a mix of good and bad is "
           "expected.", size=9, italic=True, color=META_GRAY)

    # Part C
    heading(doc, "Part C — Three short questions", before=0, new_page=True)
    for i, q in enumerate((
            "Looking at the attached graph, does the overall map reflect the "
            "learning progression you intend in this course? Why or why not?",
            "Name up to 5 important concepts and relationships you expected "
            "but don’t see.",
            "If a tool like this worked well, what would make it useful to "
            "you?"), start=1):
        p = para(doc, before=6, after=0)
        run(p, f"{i}.  {q}")
        answer_rule(doc)

    # Returning this packet
    heading(doc, "Returning this packet", before=24)
    p = para(doc, after=2)
    run(p, f"Please return by {args.return_by} to {args.return_email}.")
    p = para(doc, after=2)
    run(p, "The score tables alone are plenty; notes and Part C are welcome "
           "bonuses.")
    p = para(doc, before=14, after=0)
    run(p, "Contact: Jeffry Royce (jroyce@illinois.edu) · Prof. Alawini "
           "(alawini@illinois.edu)", size=9, color=META_GRAY)

    doc.save(args.out)
    print(f"wrote {args.out}")
    print(f"Part A: {len(head)} head + {len(tail)} tail concepts; "
          f"Part B: {len(part_b)} relationships (seed={args.seed})")


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--relations", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--course-title", required=True,
                    help='e.g. "iCAN Algorithms (CS 401/403)"')
    ap.add_argument("--run-label", required=True,
                    help='e.g. "cs401_403 combined, Qwen-14B, August 2026"')
    ap.add_argument("--banner-label", default=None,
                    help="Text after the em dash in the title; defaults to "
                         "course title")
    ap.add_argument("--sample", action="store_true",
                    help="Render the banner label in red (SAMPLE style)")
    ap.add_argument("--eval-json", default=None,
                    help="eval.py --output_json file with "
                         "node_significance/triplet_accuracy means")
    ap.add_argument("--node-significance", type=float, default=None)
    ap.add_argument("--triplet-accuracy", type=float, default=None)
    ap.add_argument("--seed", type=int, default=597)
    ap.add_argument("--head-concepts", type=int, default=12)
    ap.add_argument("--tail-concepts", type=int, default=3)
    ap.add_argument("--glance-concepts", type=int, default=8)
    ap.add_argument("--relationships", type=int, default=15)
    ap.add_argument("--justification-chars", type=int, default=420)
    ap.add_argument("--part-a-note", default=None,
                    help="Extra course-specific sentence appended to the "
                         "Part A footnote")
    ap.add_argument("--name-overrides", default=None,
                    help="JSON file mapping raw concept names to display "
                         'names, e.g. {"BIG_O_NOTATION": "Big-O notation"}')
    ap.add_argument("--return-by", default="[date]")
    ap.add_argument("--return-email", default="[email]")
    args = ap.parse_args()
    if args.banner_label is None:
        args.banner_label = args.course_title
    build(args)


if __name__ == "__main__":
    main()
